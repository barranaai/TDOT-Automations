# Dev-time tool: convert a master retainer .docx's YELLOW placeholder runs into
# docxtemplater {tags}, run-group aware, removing the highlight. Reports any
# unmapped yellow segment so nothing is silently missed.
import sys, re, os
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn

# normalized source text -> docxtemplater tag (PA master)
MAP = {
 'agreement date': '{agreementDate}', 'agreement_date': '{agreementDate}',
 'name of principal applicant': '{paName}', 'principal applicant': '{paName}', 'client': '{paName}',
 'address of the principal applicant': '{paAddress}', 'address': '{paAddress}',
 'principal applicant phone number': '{paPhone}', "principal applicant's email": '{paEmail}',
 'type of application': '{applicationType}',
 'xx': '{scopeAnnexNo}', 'x': '{paymentAnnexNo}',
 'payment_terms_summary_fees': '{serviceFees}', 'professional taxes cad': '{hst}',
 'total': '{total}', 'xxx': '{govFee}',
}
def norm(s): return re.sub(r'\s+', ' ', s).strip().lower().replace('’', "'")
def is_yellow(run):
    hc = run.font.highlight_color
    if hc is not None and hc != WD_COLOR_INDEX.AUTO and str(hc).upper().startswith('YELLOW'): return True
    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        h = rPr.find(qn('w:highlight'))
        if h is not None and (h.get(qn('w:val')) or '').lower() == 'yellow': return True
    return False
def clear_hl(run):
    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        h = rPr.find(qn('w:highlight'))
        if h is not None: rPr.remove(h)

def iter_paras(doc):
    for p in doc.paragraphs: yield p
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs: yield p

src, out = sys.argv[1], sys.argv[2]
doc = Document(src)
mapped, unmapped, skipped = [], [], 0
for p in iter_paras(doc):
    runs = p.runs; i = 0
    while i < len(runs):
        if not is_yellow(runs[i]): i += 1; continue
        j = i
        while j < len(runs) and is_yellow(runs[j]): j += 1
        seg = runs[i:j]
        combined = ''.join(r.text for r in seg)
        key = norm(combined)
        for r in seg: clear_hl(r)
        if not key:
            skipped += 1
        elif key in MAP:
            seg[0].text = MAP[key]
            for r in seg[1:]: r.text = ''
            mapped.append((combined, MAP[key]))
        else:
            unmapped.append(combined)
        i = j
doc.save(out)
print('MAPPED %d:' % len(mapped))
for c, t in mapped: print('   "%s" -> %s' % (c[:50], t))
print('SKIPPED blank-yellow segments: %d' % skipped)
print('UNMAPPED (need attention): %s' % (unmapped if unmapped else 'NONE'))
