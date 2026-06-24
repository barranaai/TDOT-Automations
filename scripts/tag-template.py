# Tag a master by ORDERED tag list (robust for templates with repeated generic
# fields). No 3rd arg = DUMP the yellow segments in order so you can build the list.
import sys, json
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
def is_yellow(run):
    hc=run.font.highlight_color
    if hc is not None and hc!=WD_COLOR_INDEX.AUTO and str(hc).upper().startswith('YELLOW'): return True
    rPr=run._element.find(qn('w:rPr'))
    if rPr is not None:
        h=rPr.find(qn('w:highlight'))
        if h is not None and (h.get(qn('w:val')) or '').lower()=='yellow': return True
    return False
def clear_hl(run):
    rPr=run._element.find(qn('w:rPr'))
    if rPr is not None:
        h=rPr.find(qn('w:highlight'))
        if h is not None: rPr.remove(h)
def iter_paras(doc):
    for p in doc.paragraphs: yield p
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs: yield p
def segments(doc):
    out=[]
    for p in iter_paras(doc):
        runs=p.runs; i=0
        while i<len(runs):
            if not is_yellow(runs[i]): i+=1; continue
            j=i
            while j<len(runs) and is_yellow(runs[j]): j+=1
            out.append((runs[i:j], ''.join(r.text for r in runs[i:j]))); i=j
    return out
src=sys.argv[1]; doc=Document(src); segs=segments(doc)
if len(sys.argv)<4:
    for idx,(seg,txt) in enumerate(segs): print('%2d | %r'%(idx,txt))
    print('TOTAL %d'%len(segs))
else:
    out=sys.argv[2]; tags=json.loads(sys.argv[3])
    assert len(tags)==len(segs), 'tags %d != segments %d'%(len(tags),len(segs))
    for (seg,_),tag in zip(segs,tags):
        for r in seg: clear_hl(r)
        if tag:
            seg[0].text=tag
            for r in seg[1:]: r.text=''
    doc.save(out); print('Applied %d → %s'%(sum(1 for t in tags if t),out))
